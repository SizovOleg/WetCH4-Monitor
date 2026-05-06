# -*- coding: utf-8 -*-
"""Сборка отдельного .docx со списком литературы (RU + EN References).

Версия 3 — соответствует официальным правилам журнала «Современные проблемы
дистанционного зондирования Земли из космоса» (ИКИ РАН), правила-оформления-2023:

- Возвращён `DOI:` префикс (вместо https://doi.org/...).
- Удалены `(дата обращения)` / `(date of access)`.
- Romanovsky 2010 удалён по решению автора → 34 записи.
- В Литература для англоязычных записей: `V.` (не `Vol.`), `No.`, `P.`, `Article`.
- Если авторов больше 4 и есть DOI → первые 3 + «et al.» (правило п. 141).
- В References для транслитерированных русских источников — пометка `(in Russian)`.
- Зоны разделены точками (не запятыми, не тире) — правило п. 144.
- Авторы в Литература — курсивом; журналы в References — курсивом.
- Автоматическая нумерация Word (через numbering.xml).

Запуск: python build_references_docx.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ---------------------------------------------------------------------------
# Литература (RU): (authors_italic, rest_plain)
# Авторы > 4 при наличии DOI: первые 3 + «et al.»
# ---------------------------------------------------------------------------
REFERENCES_RU: list[tuple[str, str]] = [
    # 1 — Лапшина (1 автор)
    ("Лапшина Е.Д.",
     "Флора болот юго-востока Западной Сибири. Томск: Изд-во Томского ун-та, "
     "2004. 295 с."),
    # 2 — Лисс (9 авторов, без DOI → все)
    ("Лисс О.Л., Абрамова Л.И., Аветов Н.А., Березина Н.А., Инишева Л.И., "
     "Курнишкова Т.В., Слука З.А., Толпышева Т.Ю., Шведчикова Н.К.",
     "Болотные системы Западной Сибири и их природоохранное значение. "
     "Тула: Гриф и К, 2001. 584 с."),
    # 3 — Национальный атлас (без авторов)
    ("Национальный атлас России.",
     "Том 2: Природа. Экология. М.: Роскартография, 2004. 495 с. "
     "(карты физико-географического районирования)"),
    # 4 — Buchhorn (6 авторов с DOI → первые 3 + et al.)
    ("Buchhorn M., Lesiv M., Tsendbazar N.-E. et al.",
     "Copernicus Global Land Cover Layers — Collection 2 // Remote Sensing. "
     "2020. V. 12. No. 6. Article 1044. DOI: 10.3390/rs12061044."),
    # 5 — Chechin (12 авторов с DOI → первые 3 + et al.)
    ("Chechin D.G., Repina I.A., Artamonov A.Yu. et al.",
     "Quantifying spatial heterogeneities of surface heat budget and methane "
     "emissions over West-Siberian peatland: Highlights from the Mukhrino 2022 "
     "campaign // Forests. 2024. V. 15. No. 1. Article 102. "
     "DOI: 10.3390/f15010102."),
    # 6 — CGLS (без авторов; URL остаётся как нативный URL продукта)
    ("Copernicus Global Land Service (CGLS).",
     "Land Cover 100 m, Collection 3, Epoch 2019, Globe [data product]. "
     "URL: https://land.copernicus.eu/global/products/lc. "
     "[Описание методики продукта см. Buchhorn et al., 2020]."),
    # 7 — Didan (1 автор + DOI)
    ("Didan K.",
     "MODIS/Terra Vegetation Indices 16-Day L3 Global 500m SIN Grid V061 "
     "[data set]. NASA EOSDIS Land Processes DAAC, 2021. "
     "DOI: 10.5067/MODIS/MOD13A1.061."),
    # 8 — Dyukarev (5 авторов с DOI → первые 3 + et al.)
    ("Dyukarev E.A., Veretennikova E.E., Sabrekov A.F. et al.",
     "Methane and carbon dioxide fluxes correlation according to automatic "
     "chamber observations at the Mukhrino bog ridge and hollow complex // "
     "Environmental Dynamics and Global Climate Change. 2024. V. 15. No. 4. "
     "P. 276–288. DOI: 10.18822/edgcc636456."),
    # 9 — Glagolev (5 авторов с DOI → первые 3 + et al.)
    ("Glagolev M., Kleptsova I., Filippov I. et al.",
     "Regional methane emission from West Siberia mire landscapes // "
     "Environmental Research Letters. 2011. V. 6. No. 4. Article 045214. "
     "DOI: 10.1088/1748-9326/6/4/045214."),
    # 10 — Gorelick (6 авторов с DOI → первые 3 + et al.)
    ("Gorelick N., Hancher M., Dixon M. et al.",
     "Google Earth Engine: Planetary-scale geospatial analysis for everyone // "
     "Remote Sensing of Environment. 2017. V. 202. P. 18–27. "
     "DOI: 10.1016/j.rse.2017.06.031."),
    # 11 — IPCC (без индивидуальных авторов)
    ("IPCC, 2021:",
     "Climate Change 2021: The Physical Science Basis. Contribution of Working "
     "Group I to the Sixth Assessment Report of the Intergovernmental Panel on "
     "Climate Change / Masson-Delmotte V., Zhai P., Pirani A. et al. (eds.). "
     "Cambridge: Cambridge University Press, 2021. 2391 p. "
     "DOI: 10.1017/9781009157896."),
    # 12 — Kim (7 авторов с DOI → первые 3 + et al.)
    ("Kim H.-S., Maksyutov S., Glagolev M.V. et al.",
     "Evaluation of methane emissions from West Siberian wetlands based on "
     "inverse modeling // Environmental Research Letters. 2011. V. 6. No. 3. "
     "Article 035201. DOI: 10.1088/1748-9326/6/3/035201."),
    # 13 — Knox (3 явных + et al.)
    ("Knox S.H., Bansal S., McNicol G. et al.",
     "Identifying dominant environmental predictors of freshwater wetland "
     "methane fluxes across diurnal to seasonal time scales // Global Change "
     "Biology. 2021. V. 27. No. 15. P. 3582–3604. DOI: 10.1111/gcb.15661."),
    # 14 — Lan (3 автора)
    ("Lan X., Thoning K.W., Dlugokencky E.J.",
     "Trends in globally-averaged CH₄, N₂O, and SF₆ determined from NOAA "
     "Global Monitoring Laboratory measurements, Version 2024-08. "
     "NOAA Global Monitoring Laboratory, 2024. DOI: 10.15138/P8XG-AA10."),
    # 15 — Lindqvist (11 авторов с DOI → первые 3 + et al.)
    ("Lindqvist H., Kivimäki E., Häkkilä T. et al.",
     "Evaluation of Sentinel-5P TROPOMI Methane Observations at Northern High "
     "Latitudes // Remote Sensing. 2024. V. 16. No. 16. Article 2979. "
     "DOI: 10.3390/rs16162979."),
    # 16 — Lorente (3 явных + et al.)
    ("Lorente A., Borsdorff T., Butz A. et al.",
     "Methane retrieved from TROPOMI: improvement of the data product and "
     "validation of the first two years of measurements // Atmospheric "
     "Measurement Techniques. 2021. V. 14. No. 1. P. 665–684. "
     "DOI: 10.5194/amt-14-665-2021."),
    # 17 — Mastepanov (7 авторов с DOI → первые 3 + et al.)
    ("Mastepanov M., Sigsgaard C., Dlugokencky E.J. et al.",
     "Large tundra methane burst during onset of freezing // Nature. 2008. "
     "V. 456. No. 7222. P. 628–630. DOI: 10.1038/nature07464."),
    # 18 — McNicol (3 явных + et al.)
    ("McNicol G., Fluet-Chouinard E., Ouyang Z. et al.",
     "Upscaling wetland methane emissions from the FLUXNET-CH4 eddy covariance "
     "network (UpCH4 v1.0): Model development, network assessment, and budget "
     "comparison // AGU Advances. 2023. V. 4. No. 5. Article e2023AV000956. "
     "DOI: 10.1029/2023AV000956."),
    # 19 — Muñoz Sabater (3 явных + et al.)
    ("Muñoz Sabater J., Dutra E., Agustí-Panareda A. et al.",
     "ERA5-Land: a state-of-the-art global reanalysis dataset for land "
     "applications // Earth System Science Data. 2021. V. 13. No. 9. "
     "P. 4349–4383. DOI: 10.5194/essd-13-4349-2021."),
    # 20 — Panikov (2 автора)
    ("Panikov N.S., Dedysh S.N.",
     "Cold season CH₄ and CO₂ emission from boreal peat bogs (West Siberia): "
     "Winter fluxes and thaw activation dynamics // Global Biogeochemical "
     "Cycles. 2000. V. 14. No. 4. P. 1071–1080. DOI: 10.1029/1999GB900097."),
    # 21 — Rinne (3 явных + et al.)
    ("Rinne J., Tuittila E.-S., Peltola O. et al.",
     "Temporal variation of ecosystem scale methane emission from a boreal fen "
     "in relation to temperature, water table position, and carbon dioxide "
     "fluxes // Global Biogeochemical Cycles. 2018. V. 32. No. 7. P. 1087–1106. "
     "DOI: 10.1029/2017GB005747."),
    # 22 — Sabrekov 2013 (5 авторов с DOI → первые 3 + et al.)
    ("Sabrekov A.F., Glagolev M.V., Kleptsova I.E. et al.",
     "Methane emission from mires of the West Siberian taiga // Eurasian Soil "
     "Science. 2013. V. 46. No. 12. P. 1182–1193. "
     "DOI: 10.1134/S1064229314010098."),
    # 23 — Sabrekov 2011 (5 авторов, без DOI → все)
    ("Sabrekov A.F., Kleptsova I.E., Glagolev M.V., Maksyutov S.S., Machida T.",
     "Methane emission from middle taiga oligotrophic hollows of Western "
     "Siberia // Tomsk State Pedagogical University Bulletin. 2011. "
     "No. 5 (107). P. 135–143."),
    # 24 — Saunois (3 явных + et al.)
    ("Saunois M., Stavert A.R., Poulter B. et al.",
     "The Global Methane Budget 2000–2017 // Earth System Science Data. 2020. "
     "V. 12. No. 3. P. 1561–1623. DOI: 10.5194/essd-12-1561-2020."),
    # 25 — Segers (1 автор)
    ("Segers R.",
     "Methane production and methane consumption: a review of processes "
     "underlying wetland methane fluxes // Biogeochemistry. 1998. V. 41. "
     "No. 1. P. 23–51. DOI: 10.1023/A:1005929032764."),
    # 26 — Sheng (3 явных + et al.)
    ("Sheng Y., Smith L.C., MacDonald G.M. et al.",
     "A high-resolution GIS-based inventory of the west Siberian peat carbon "
     "pool // Global Biogeochemical Cycles. 2004. V. 18. No. 3. "
     "Article GB3004. DOI: 10.1029/2003GB002190."),
    # 27 — Terentieva (5 авторов с DOI → первые 3 + et al.)
    ("Terentieva I.E., Glagolev M.V., Lapshina E.D. et al.",
     "Mapping of West Siberian taiga wetland complexes using Landsat imagery: "
     "implications for methane emissions // Biogeosciences. 2016. V. 13. "
     "No. 16. P. 4615–4626. DOI: 10.5194/bg-13-4615-2016."),
    # 28 — Tsuruta (3 явных + et al.)
    ("Tsuruta A., Aalto T., Backman L. et al.",
     "Methane budget estimates in Finland from the CarbonTracker Europe-CH₄ "
     "data assimilation system // Tellus B. 2019. V. 71. No. 1. "
     "Article 1565030. DOI: 10.1080/16000889.2018.1565030."),
    # 29 — Veretennikova (2 автора)
    ("Veretennikova E.E., Dyukarev E.A.",
     "Comparison of methane fluxes of open and forested bogs of the southern "
     "taiga zone of Western Siberia // Boreal Environment Research. 2021. "
     "V. 26. P. 43–59."),
    # 30 — Winderlich (4 автора — НЕ больше 4, оставляем все)
    ("Winderlich J., Gerbig C., Kolle O., Heimann M.",
     "Inferences from CO₂ and CH₄ concentration profiles at the Zotino Tall "
     "Tower Observatory (ZOTTO) on regional summertime ecosystem fluxes // "
     "Biogeosciences. 2014. V. 11. No. 7. P. 2055–2068. "
     "DOI: 10.5194/bg-11-2055-2014."),
    # 31 — Xu (4 автора — оставляем все)
    ("Xu J., Morris P.J., Liu J., Holden J.",
     "PEATMAP: Refining estimates of global peatland distribution based on a "
     "meta-analysis // CATENA. 2018. V. 160. P. 134–140. "
     "DOI: 10.1016/j.catena.2017.09.010."),
    # 32 — Ying (10 авторов с DOI → первые 3 + et al.)
    ("Ying Q., Poulter B., Watts J.D. et al.",
     "WetCH4: a machine-learning-based upscaling of methane fluxes of northern "
     "wetlands during 2016–2022 // Earth System Science Data. 2025. V. 17. "
     "No. 6. P. 2507–2534. DOI: 10.5194/essd-17-2507-2025."),
    # 33 — Yuan (3 явных + et al.)
    ("Yuan K., Zhu Q., Riley W.J. et al.",
     "Boreal-Arctic wetland methane emissions modulated by warming and "
     "vegetation activity // Nature Climate Change. 2024. V. 14. No. 3. "
     "P. 282–288. DOI: 10.1038/s41558-024-01933-3."),
    # 34 — Yvon-Durocher (8 авторов с DOI → первые 3 + et al.)
    ("Yvon-Durocher G., Allen A.P., Bastviken D. et al.",
     "Methane fluxes show consistent temperature dependence across microbial "
     "to ecosystem scales // Nature. 2014. V. 507. No. 7493. P. 488–491. "
     "DOI: 10.1038/nature13164."),
]


# ---------------------------------------------------------------------------
# References (EN): (authors_plain, before_journal, journal_italic, after_journal)
# Разделители — запятые, журналы курсивом, для русских источников — (in Russian).
# ---------------------------------------------------------------------------
REFERENCES_EN: list[tuple[str, str, str, str]] = [
    # 1 — Лапшина (book, Russian original)
    ("Lapshina E.D., ",
     "Flora bolot yugo-vostoka Zapadnoi Sibiri (Bog flora of the south-east of "
     "Western Siberia), Tomsk: Tomsk University Press, 2004, 295 p. "
     "(in Russian).",
     "", ""),
    # 2 — Лисс (book, Russian original)
    ("Liss O.L., Abramova L.I., Avetov N.A., Berezina N.A., Inisheva L.I., "
     "Kurnishkova T.V., Sluka Z.A., Tolpysheva T.Yu., Shvedchikova N.K., ",
     "Bolotnye sistemy Zapadnoi Sibiri i ikh prirodookhrannoe znachenie "
     "(Wetland systems of Western Siberia and their environmental "
     "importance), Tula: Grif i K, 2001, 584 p. (in Russian).",
     "", ""),
    # 3 — Национальный атлас (Russian)
    ("",
     "Natsional'nyi atlas Rossii (National Atlas of Russia), V. 2: Priroda. "
     "Ekologiya (Nature. Ecology), Moscow: Roskartografiya, 2004, 495 p. "
     "(in Russian).",
     "", ""),
    # 4 — Buchhorn
    ("Buchhorn M., Lesiv M., Tsendbazar N.-E. et al., Copernicus Global Land "
     "Cover Layers — Collection 2, ",
     "", "Remote Sensing",
     ", 2020, V. 12, No. 6, Article 1044. DOI: 10.3390/rs12061044."),
    # 5 — Chechin
    ("Chechin D.G., Repina I.A., Artamonov A.Yu. et al., Quantifying spatial "
     "heterogeneities of surface heat budget and methane emissions over "
     "West-Siberian peatland: Highlights from the Mukhrino 2022 campaign, ",
     "", "Forests",
     ", 2024, V. 15, No. 1, Article 102. DOI: 10.3390/f15010102."),
    # 6 — CGLS (data product, no authors)
    ("",
     "Copernicus Global Land Service (CGLS), Land Cover 100 m, Collection 3, "
     "Epoch 2019, Globe [data product]. "
     "URL: https://land.copernicus.eu/global/products/lc.",
     "", ""),
    # 7 — Didan (data set)
    ("Didan K., MODIS/Terra Vegetation Indices 16-Day L3 Global 500m SIN Grid "
     "V061 [data set], NASA EOSDIS Land Processes DAAC, 2021. "
     "DOI: 10.5067/MODIS/MOD13A1.061.",
     "", "", ""),
    # 8 — Dyukarev (англоязычная статья в EDGCC)
    ("Dyukarev E.A., Veretennikova E.E., Sabrekov A.F. et al., Methane and "
     "carbon dioxide fluxes correlation according to automatic chamber "
     "observations at the Mukhrino bog ridge and hollow complex, ",
     "", "Environmental Dynamics and Global Climate Change",
     ", 2024, V. 15, No. 4, pp. 276–288. "
     "DOI: 10.18822/edgcc636456."),
    # 9 — Glagolev
    ("Glagolev M., Kleptsova I., Filippov I. et al., Regional methane "
     "emission from West Siberia mire landscapes, ",
     "", "Environmental Research Letters",
     ", 2011, V. 6, No. 4, Article 045214. "
     "DOI: 10.1088/1748-9326/6/4/045214."),
    # 10 — Gorelick
    ("Gorelick N., Hancher M., Dixon M. et al., Google Earth Engine: "
     "Planetary-scale geospatial analysis for everyone, ",
     "", "Remote Sensing of Environment",
     ", 2017, V. 202, pp. 18–27. DOI: 10.1016/j.rse.2017.06.031."),
    # 11 — IPCC
    ("IPCC, ",
     "Climate Change 2021: The Physical Science Basis. Contribution of Working "
     "Group I to the Sixth Assessment Report of the Intergovernmental Panel on "
     "Climate Change, Masson-Delmotte V., Zhai P., Pirani A. et al. (eds.), "
     "Cambridge: Cambridge University Press, 2021, 2391 p. "
     "DOI: 10.1017/9781009157896.",
     "", ""),
    # 12 — Kim
    ("Kim H.-S., Maksyutov S., Glagolev M.V. et al., Evaluation of methane "
     "emissions from West Siberian wetlands based on inverse modeling, ",
     "", "Environmental Research Letters",
     ", 2011, V. 6, No. 3, Article 035201. "
     "DOI: 10.1088/1748-9326/6/3/035201."),
    # 13 — Knox
    ("Knox S.H., Bansal S., McNicol G. et al., Identifying dominant "
     "environmental predictors of freshwater wetland methane fluxes across "
     "diurnal to seasonal time scales, ",
     "", "Global Change Biology",
     ", 2021, V. 27, No. 15, pp. 3582–3604. DOI: 10.1111/gcb.15661."),
    # 14 — Lan (data product, NOAA)
    ("Lan X., Thoning K.W., Dlugokencky E.J., Trends in globally-averaged "
     "CH₄, N₂O, and SF₆ determined from NOAA Global Monitoring Laboratory "
     "measurements, Version 2024-08, NOAA Global Monitoring Laboratory, 2024. "
     "DOI: 10.15138/P8XG-AA10.",
     "", "", ""),
    # 15 — Lindqvist
    ("Lindqvist H., Kivimäki E., Häkkilä T. et al., Evaluation of Sentinel-5P "
     "TROPOMI Methane Observations at Northern High Latitudes, ",
     "", "Remote Sensing",
     ", 2024, V. 16, No. 16, Article 2979. DOI: 10.3390/rs16162979."),
    # 16 — Lorente
    ("Lorente A., Borsdorff T., Butz A. et al., Methane retrieved from "
     "TROPOMI: improvement of the data product and validation of the first "
     "two years of measurements, ",
     "", "Atmospheric Measurement Techniques",
     ", 2021, V. 14, No. 1, pp. 665–684. DOI: 10.5194/amt-14-665-2021."),
    # 17 — Mastepanov
    ("Mastepanov M., Sigsgaard C., Dlugokencky E.J. et al., Large tundra "
     "methane burst during onset of freezing, ",
     "", "Nature",
     ", 2008, V. 456, No. 7222, pp. 628–630. DOI: 10.1038/nature07464."),
    # 18 — McNicol
    ("McNicol G., Fluet-Chouinard E., Ouyang Z. et al., Upscaling wetland "
     "methane emissions from the FLUXNET-CH4 eddy covariance network "
     "(UpCH4 v1.0): Model development, network assessment, and budget "
     "comparison, ",
     "", "AGU Advances",
     ", 2023, V. 4, No. 5, Article e2023AV000956. "
     "DOI: 10.1029/2023AV000956."),
    # 19 — Muñoz Sabater
    ("Muñoz Sabater J., Dutra E., Agustí-Panareda A. et al., ERA5-Land: a "
     "state-of-the-art global reanalysis dataset for land applications, ",
     "", "Earth System Science Data",
     ", 2021, V. 13, No. 9, pp. 4349–4383. "
     "DOI: 10.5194/essd-13-4349-2021."),
    # 20 — Panikov
    ("Panikov N.S., Dedysh S.N., Cold season CH₄ and CO₂ emission from boreal "
     "peat bogs (West Siberia): Winter fluxes and thaw activation dynamics, ",
     "", "Global Biogeochemical Cycles",
     ", 2000, V. 14, No. 4, pp. 1071–1080. DOI: 10.1029/1999GB900097."),
    # 21 — Rinne
    ("Rinne J., Tuittila E.-S., Peltola O. et al., Temporal variation of "
     "ecosystem scale methane emission from a boreal fen in relation to "
     "temperature, water table position, and carbon dioxide fluxes, ",
     "", "Global Biogeochemical Cycles",
     ", 2018, V. 32, No. 7, pp. 1087–1106. DOI: 10.1029/2017GB005747."),
    # 22 — Sabrekov 2013 (English translation в Eurasian Soil Science)
    ("Sabrekov A.F., Glagolev M.V., Kleptsova I.E. et al., Methane emission "
     "from mires of the West Siberian taiga, ",
     "", "Eurasian Soil Science",
     ", 2013, V. 46, No. 12, pp. 1182–1193. "
     "DOI: 10.1134/S1064229314010098."),
    # 23 — Sabrekov 2011 (Russian original)
    ("Sabrekov A.F., Kleptsova I.E., Glagolev M.V., Maksyutov S.S., "
     "Machida T., Methane emission from middle taiga oligotrophic hollows of "
     "Western Siberia, ",
     "", "Tomsk State Pedagogical University Bulletin",
     ", 2011, No. 5 (107), pp. 135–143 (in Russian)."),
    # 24 — Saunois
    ("Saunois M., Stavert A.R., Poulter B. et al., The Global Methane Budget "
     "2000–2017, ",
     "", "Earth System Science Data",
     ", 2020, V. 12, No. 3, pp. 1561–1623. "
     "DOI: 10.5194/essd-12-1561-2020."),
    # 25 — Segers
    ("Segers R., Methane production and methane consumption: a review of "
     "processes underlying wetland methane fluxes, ",
     "", "Biogeochemistry",
     ", 1998, V. 41, No. 1, pp. 23–51. DOI: 10.1023/A:1005929032764."),
    # 26 — Sheng
    ("Sheng Y., Smith L.C., MacDonald G.M. et al., A high-resolution GIS-based "
     "inventory of the west Siberian peat carbon pool, ",
     "", "Global Biogeochemical Cycles",
     ", 2004, V. 18, No. 3, Article GB3004. DOI: 10.1029/2003GB002190."),
    # 27 — Terentieva
    ("Terentieva I.E., Glagolev M.V., Lapshina E.D. et al., Mapping of West "
     "Siberian taiga wetland complexes using Landsat imagery: implications "
     "for methane emissions, ",
     "", "Biogeosciences",
     ", 2016, V. 13, No. 16, pp. 4615–4626. DOI: 10.5194/bg-13-4615-2016."),
    # 28 — Tsuruta
    ("Tsuruta A., Aalto T., Backman L. et al., Methane budget estimates in "
     "Finland from the CarbonTracker Europe-CH₄ data assimilation system, ",
     "", "Tellus B",
     ", 2019, V. 71, No. 1, Article 1565030. "
     "DOI: 10.1080/16000889.2018.1565030."),
    # 29 — Veretennikova
    ("Veretennikova E.E., Dyukarev E.A., Comparison of methane fluxes of open "
     "and forested bogs of the southern taiga zone of Western Siberia, ",
     "", "Boreal Environment Research",
     ", 2021, V. 26, pp. 43–59."),
    # 30 — Winderlich (4 авторов — оставляем всех)
    ("Winderlich J., Gerbig C., Kolle O., Heimann M., Inferences from CO₂ and "
     "CH₄ concentration profiles at the Zotino Tall Tower Observatory (ZOTTO) "
     "on regional summertime ecosystem fluxes, ",
     "", "Biogeosciences",
     ", 2014, V. 11, No. 7, pp. 2055–2068. DOI: 10.5194/bg-11-2055-2014."),
    # 31 — Xu (4 авторов — оставляем всех)
    ("Xu J., Morris P.J., Liu J., Holden J., PEATMAP: Refining estimates of "
     "global peatland distribution based on a meta-analysis, ",
     "", "CATENA",
     ", 2018, V. 160, pp. 134–140. DOI: 10.1016/j.catena.2017.09.010."),
    # 32 — Ying
    ("Ying Q., Poulter B., Watts J.D. et al., WetCH4: a machine-learning-based "
     "upscaling of methane fluxes of northern wetlands during 2016–2022, ",
     "", "Earth System Science Data",
     ", 2025, V. 17, No. 6, pp. 2507–2534. "
     "DOI: 10.5194/essd-17-2507-2025."),
    # 33 — Yuan
    ("Yuan K., Zhu Q., Riley W.J. et al., Boreal-Arctic wetland methane "
     "emissions modulated by warming and vegetation activity, ",
     "", "Nature Climate Change",
     ", 2024, V. 14, No. 3, pp. 282–288. DOI: 10.1038/s41558-024-01933-3."),
    # 34 — Yvon-Durocher
    ("Yvon-Durocher G., Allen A.P., Bastviken D. et al., Methane fluxes show "
     "consistent temperature dependence across microbial to ecosystem "
     "scales, ",
     "", "Nature",
     ", 2014, V. 507, No. 7493, pp. 488–491. DOI: 10.1038/nature13164."),
]


# ---------------------------------------------------------------------------
# DOCX-сборка
# ---------------------------------------------------------------------------

def add_heading(doc: Document, text: str) -> None:
    """Заголовок раздела (центрированный, жирный)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)


def setup_numbering(doc: Document) -> tuple[int, int]:
    """Создать две независимые автонумерации (RU/EN) в numbering.xml."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    numbering = doc.part.numbering_part.element

    def make_abstract(abs_id: int) -> "OxmlElement":
        a = OxmlElement("w:abstractNum")
        a.set(qn("w:abstractNumId"), str(abs_id))
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        s = OxmlElement("w:start")
        s.set(qn("w:val"), "1")
        lvl.append(s)
        nf = OxmlElement("w:numFmt")
        nf.set(qn("w:val"), "decimal")
        lvl.append(nf)
        lt = OxmlElement("w:lvlText")
        lt.set(qn("w:val"), "%1.")
        lvl.append(lt)
        lj = OxmlElement("w:lvlJc")
        lj.set(qn("w:val"), "left")
        lvl.append(lj)
        ppr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "567")
        ind.set(qn("w:hanging"), "567")
        ppr.append(ind)
        lvl.append(ppr)
        a.append(lvl)
        return a

    def make_num(num_id: int, abs_id: int) -> "OxmlElement":
        n = OxmlElement("w:num")
        n.set(qn("w:numId"), str(num_id))
        ai = OxmlElement("w:abstractNumId")
        ai.set(qn("w:val"), str(abs_id))
        n.append(ai)
        return n

    existing_abs = [
        int(a.get(qn("w:abstractNumId")))
        for a in numbering.findall(qn("w:abstractNum"))
    ]
    existing_num = [
        int(n.get(qn("w:numId")))
        for n in numbering.findall(qn("w:num"))
    ]
    abs_id_ru = max(existing_abs, default=-1) + 1
    abs_id_en = abs_id_ru + 1
    num_id_ru = max(existing_num, default=0) + 1
    num_id_en = num_id_ru + 1

    abstr_ru = make_abstract(abs_id_ru)
    abstr_en = make_abstract(abs_id_en)
    first_num = numbering.find(qn("w:num"))
    if first_num is not None:
        first_num.addprevious(abstr_ru)
        first_num.addprevious(abstr_en)
    else:
        numbering.append(abstr_ru)
        numbering.append(abstr_en)

    numbering.append(make_num(num_id_ru, abs_id_ru))
    numbering.append(make_num(num_id_en, abs_id_en))

    return num_id_ru, num_id_en


def ensure_numbering_part(doc: Document) -> None:
    """Гарантировать наличие numbering part."""
    if doc.part.numbering_part is not None:
        return
    p = doc.add_paragraph(style="List Number")
    p.text = "PLACEHOLDER"
    p._element.getparent().remove(p._element)


def apply_num_id(paragraph, num_id: int) -> None:
    """Применить numId к абзацу через w:numPr."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    pPr = paragraph._element.get_or_add_pPr()
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        numPr = OxmlElement("w:numPr")
        pPr.append(numPr)
    else:
        for child in list(numPr):
            numPr.remove(child)

    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numPr.append(ilvl)

    n = OxmlElement("w:numId")
    n.set(qn("w:val"), str(num_id))
    numPr.append(n)


def add_ru_entry(doc: Document, num_id: int, authors_italic: str,
                 rest_plain: str) -> None:
    """Запись в Литература (авторы курсивом)."""
    p = doc.add_paragraph()
    apply_num_id(p, num_id)

    if authors_italic:
        run_auth = p.add_run(authors_italic)
        run_auth.italic = True
        run_auth.font.size = Pt(11)
        run_space = p.add_run(" ")
        run_space.font.size = Pt(11)

    run_rest = p.add_run(rest_plain)
    run_rest.font.size = Pt(11)


def add_en_entry(doc: Document, num_id: int, authors_plain: str,
                 before_journal: str, journal_italic: str,
                 after_journal: str) -> None:
    """Запись в References (журнал курсивом)."""
    p = doc.add_paragraph()
    apply_num_id(p, num_id)

    if authors_plain:
        run_auth = p.add_run(authors_plain)
        run_auth.font.size = Pt(11)
    if before_journal:
        run_before = p.add_run(before_journal)
        run_before.font.size = Pt(11)
    if journal_italic:
        run_journal = p.add_run(journal_italic)
        run_journal.italic = True
        run_journal.font.size = Pt(11)
    if after_journal:
        run_after = p.add_run(after_journal)
        run_after.font.size = Pt(11)


def main() -> None:
    """Сборка финального .docx."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    ensure_numbering_part(doc)
    num_id_ru, num_id_en = setup_numbering(doc)

    add_heading(doc, "Литература")
    for authors, rest in REFERENCES_RU:
        add_ru_entry(doc, num_id_ru, authors, rest)

    doc.add_page_break()

    add_heading(doc, "References")
    for authors, before, journal, after in REFERENCES_EN:
        add_en_entry(doc, num_id_en, authors, before, journal, after)

    out_path = Path(__file__).parent / "REFERENCES_final.docx"
    try:
        doc.save(out_path)
    except PermissionError:
        out_path = Path(__file__).parent / "REFERENCES_final_v2.docx"
        doc.save(out_path)
        print("⚠️  Исходный файл занят (открыт в Word) — сохранено с суффиксом v2.")
    print(f"Saved: {out_path}")
    print(f"Литература: {len(REFERENCES_RU)} entries (numId={num_id_ru})")
    print(f"References:  {len(REFERENCES_EN)} entries (numId={num_id_en})")


if __name__ == "__main__":
    import sys
    sys.stdout.reconfigure(encoding="utf-8")
    main()
